from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Header, Query, Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone
import requests
from emergentintegrations.llm.chat import LlmChat, UserMessage
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Preformatted
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import re
import xml.etree.ElementTree as ET

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Object Storage Setup
STORAGE_URL = "https://integrations.emergentagent.com/objstore/api/v1/storage"
EMERGENT_KEY = os.environ.get("EMERGENT_LLM_KEY")
APP_NAME = "jenkins-doc-gen"
storage_key = None

def init_storage():
    global storage_key
    if storage_key:
        return storage_key
    resp = requests.post(f"{STORAGE_URL}/init", json={"emergent_key": EMERGENT_KEY}, timeout=30)
    resp.raise_for_status()
    storage_key = resp.json()["storage_key"]
    return storage_key

def put_object(path: str, data: bytes, content_type: str) -> dict:
    key = init_storage()
    resp = requests.put(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key, "Content-Type": content_type},
        data=data, timeout=120
    )
    resp.raise_for_status()
    return resp.json()

def get_object(path: str) -> tuple:
    key = init_storage()
    resp = requests.get(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key}, timeout=60
    )
    resp.raise_for_status()
    return resp.content, resp.headers.get("Content-Type", "application/octet-stream")

app = FastAPI()
api_router = APIRouter(prefix="/api")

# Pydantic Models
class FileUploadResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    storage_path: str
    original_filename: str
    content_type: str
    size: int
    created_at: str

class XmlUploadResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    storage_path: str
    original_filename: str
    content_type: str
    size: int
    parsed_data: dict
    created_at: str

class DocumentationCreate(BaseModel):
    title: str
    file_ids: List[str] = []
    xml_file_ids: List[str] = []
    template_id: Optional[str] = None
    ai_provider: str = "claude"

class DocumentationUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None

class Documentation(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    title: str
    content: str
    file_ids: List[str]
    template_id: Optional[str] = None
    created_at: str
    updated_at: str

class TemplateCreate(BaseModel):
    name: str
    description: str
    prompt_template: str

class Template(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    description: str
    prompt_template: str
    created_at: str

# API Routes
@api_router.get("/")
async def root():
    return {"message": "Jenkins Documentation Generator API"}

def parse_jenkins_xml(xml_content: str) -> dict:
    """Parse Jenkins config.xml and extract key fields."""
    parsed = {
        "job_type": "unknown",
        "description": "",
        "parameters": [],
        "scm": {},
        "triggers": [],
        "build_steps": [],
        "post_build_actions": [],
        "environment_vars": [],
        "pipeline_script": "",
        "agent_label": "",
        "properties": [],
    }
    try:
        # Fix XML 1.1 declaration (Python's ET only supports 1.0)
        cleaned = re.sub(r"<\?xml[^?]*\?>", '<?xml version="1.0" encoding="UTF-8"?>', xml_content)
        root = ET.fromstring(cleaned)
        parsed["job_type"] = root.tag

        desc = root.find("description")
        if desc is not None and desc.text:
            parsed["description"] = desc.text.strip()

        # Agent / node label
        for tag in ["assignedNode", "label"]:
            node = root.find(tag)
            if node is not None and node.text:
                parsed["agent_label"] = node.text.strip()
                break

        # Parameters
        for param_def in root.iter():
            if param_def.tag.endswith("ParameterDefinition"):
                p = {}
                name_el = param_def.find("name")
                if name_el is not None and name_el.text:
                    p["name"] = name_el.text
                default_el = param_def.find("defaultValue")
                if default_el is not None and default_el.text:
                    p["default"] = default_el.text
                desc_el = param_def.find("description")
                if desc_el is not None and desc_el.text:
                    p["description"] = desc_el.text
                p["type"] = param_def.tag.replace("ParameterDefinition", "").split(".")[-1]
                if p.get("name"):
                    parsed["parameters"].append(p)

        # SCM (Git)
        for scm in root.iter("scm"):
            scm_class = scm.get("class", "")
            if "git" in scm_class.lower() or "Git" in scm_class:
                urls = [u.text for u in scm.iter("url") if u.text]
                branches = [b.text for b in scm.iter("name") if b.text and "/" in b.text]
                parsed["scm"] = {"type": "git", "urls": urls, "branches": branches}
                break

        # Triggers
        for trigger in root.iter("triggers"):
            for child in trigger:
                t = {"type": child.tag}
                spec = child.find("spec")
                if spec is not None and spec.text:
                    t["schedule"] = spec.text.strip()
                parsed["triggers"].append(t)

        # Build steps (shell, batch, etc.)
        for builders in root.iter("builders"):
            for step in builders:
                s = {"type": step.tag}
                cmd = step.find("command")
                if cmd is not None and cmd.text:
                    s["command"] = cmd.text.strip()[:500]
                script_el = step.find("scriptText")
                if script_el is not None and script_el.text:
                    s["script"] = script_el.text.strip()[:500]
                parsed["build_steps"].append(s)

        # Pipeline script (workflow/pipeline jobs)
        for script_el in root.iter("script"):
            if script_el.text and len(script_el.text.strip()) > 10:
                parsed["pipeline_script"] = script_el.text.strip()
                break

        # Script path for pipeline from SCM
        script_path = root.find(".//scriptPath")
        if script_path is not None and script_path.text:
            parsed["pipeline_script"] = f"Jenkinsfile path: {script_path.text}"

        # Post-build actions
        for publishers in root.iter("publishers"):
            for action in publishers:
                a = {"type": action.tag}
                parsed["post_build_actions"].append(a)

        # Environment variables
        for env_el in root.iter("EnvInjectJobPropertyInfo"):
            content = env_el.find("propertiesContent")
            if content is not None and content.text:
                for line in content.text.strip().split("\n"):
                    if "=" in line:
                        parsed["environment_vars"].append(line.strip())
        for env_el in root.iter("envVars"):
            for tree_map in env_el.iter():
                if tree_map.tag == "string":
                    parsed["environment_vars"].append(tree_map.text or "")

    except ET.ParseError as e:
        parsed["parse_error"] = str(e)
        # Fallback: try to extract key info with regex
        try:
            pipeline_match = re.search(r'<script>(.*?)</script>', xml_content, re.DOTALL)
            if pipeline_match:
                parsed["pipeline_script"] = pipeline_match.group(1).strip()
            desc_match = re.search(r'<description>(.*?)</description>', xml_content, re.DOTALL)
            if desc_match:
                parsed["description"] = desc_match.group(1).strip()
            for m in re.finditer(r'<name>([\w_-]+)</name>\s*<description>(.*?)</description>', xml_content, re.DOTALL):
                parsed["parameters"].append({"name": m.group(1), "description": m.group(2).strip(), "type": "String"})
            # Detect job type from root tag
            root_match = re.match(r'<\?xml[^?]*\?>\s*<(\S+)', xml_content)
            if root_match:
                parsed["job_type"] = root_match.group(1).split()[0]
        except Exception:
            pass
    return parsed

@api_router.post("/upload", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    try:
        ext = file.filename.split(".")[-1] if "." in file.filename else "bin"
        path = f"{APP_NAME}/uploads/{uuid.uuid4()}.{ext}"
        data = await file.read()
        result = put_object(path, data, file.content_type or "application/octet-stream")
        
        file_doc = {
            "id": str(uuid.uuid4()),
            "storage_path": result["path"],
            "original_filename": file.filename,
            "content_type": file.content_type,
            "size": result["size"],
            "is_deleted": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.files.insert_one(file_doc)
        return FileUploadResponse(**file_doc)
    except Exception as e:
        logging.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/upload-xml", response_model=XmlUploadResponse)
async def upload_xml(file: UploadFile = File(...)):
    try:
        data = await file.read()
        xml_content = data.decode("utf-8", errors="replace")
        parsed_data = parse_jenkins_xml(xml_content)

        ext = file.filename.split(".")[-1] if "." in file.filename else "xml"
        path = f"{APP_NAME}/xml/{uuid.uuid4()}.{ext}"
        result = put_object(path, data, "application/xml")

        xml_doc = {
            "id": str(uuid.uuid4()),
            "storage_path": result["path"],
            "original_filename": file.filename,
            "content_type": "application/xml",
            "size": result["size"],
            "raw_xml": xml_content,
            "parsed_data": parsed_data,
            "is_deleted": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.xml_files.insert_one(xml_doc)
        return XmlUploadResponse(
            id=xml_doc["id"],
            storage_path=xml_doc["storage_path"],
            original_filename=xml_doc["original_filename"],
            content_type=xml_doc["content_type"],
            size=xml_doc["size"],
            parsed_data=xml_doc["parsed_data"],
            created_at=xml_doc["created_at"]
        )
    except Exception as e:
        logging.error(f"XML upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/xml-files/{xml_id}")
async def get_xml_file(xml_id: str):
    record = await db.xml_files.find_one({"id": xml_id, "is_deleted": False}, {"_id": 0, "raw_xml": 0})
    if not record:
        raise HTTPException(status_code=404, detail="XML file not found")
    return record

@api_router.get("/files/{file_id}")
async def get_file(file_id: str, authorization: str = Header(None), auth: str = Query(None)):
    try:
        record = await db.files.find_one({"id": file_id, "is_deleted": False}, {"_id": 0})
        if not record:
            raise HTTPException(status_code=404, detail="File not found")
        
        data, content_type = get_object(record["storage_path"])
        return Response(content=data, media_type=record.get("content_type", content_type))
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get file error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/generate", response_model=Documentation)
async def generate_documentation(doc_create: DocumentationCreate):
    try:
        has_images = len(doc_create.file_ids) > 0
        has_xml = len(doc_create.xml_file_ids) > 0

        if not has_images and not has_xml:
            raise HTTPException(status_code=400, detail="Please provide screenshots or XML config files")

        # Get template if provided
        template_prompt = ""
        if doc_create.template_id:
            template = await db.templates.find_one({"id": doc_create.template_id}, {"_id": 0})
            if template:
                template_prompt = template["prompt_template"]

        # Build system prompt
        base_prompt = template_prompt if template_prompt else """You are an expert DevOps documentation specialist with deep knowledge of Jenkins pipelines, CI/CD, and infrastructure automation.

Analyze the provided Jenkins pipeline information and create comprehensive, professional documentation in Markdown format.

Focus on:
- Pipeline structure and stages
- Job parameters and their purposes
- SCM configuration (Git repos, branches)
- Build steps and shell commands
- Environment variables and configurations
- Triggers and scheduling
- Post-build actions and notifications
- Deployment processes
- Error handling and best practices

Provide clear, well-structured documentation with proper headings, code blocks (use ```groovy or ```bash for Jenkins code), and explanations."""

        # Build user message parts
        message_parts = []
        message_parts.append(f"Create detailed documentation for: {doc_create.title}\n")

        # Process XML config files
        if has_xml:
            xml_files = await db.xml_files.find(
                {"id": {"$in": doc_create.xml_file_ids}, "is_deleted": False},
                {"_id": 0}
            ).to_list(100)

            for i, xf in enumerate(xml_files):
                parsed = xf.get("parsed_data", {})
                raw_xml = xf.get("raw_xml", "")

                # Pre-parsed structured summary
                message_parts.append(f"\n--- Jenkins Job Config #{i+1}: {xf['original_filename']} ---")
                message_parts.append(f"Job Type: {parsed.get('job_type', 'unknown')}")
                if parsed.get("description"):
                    message_parts.append(f"Description: {parsed['description']}")
                if parsed.get("agent_label"):
                    message_parts.append(f"Agent/Node: {parsed['agent_label']}")
                if parsed.get("parameters"):
                    message_parts.append("Parameters:")
                    for p in parsed["parameters"]:
                        message_parts.append(f"  - {p.get('name', '?')} ({p.get('type', '?')}): {p.get('description', 'N/A')} [default: {p.get('default', 'N/A')}]")
                if parsed.get("scm", {}).get("urls"):
                    message_parts.append(f"SCM URLs: {', '.join(parsed['scm']['urls'])}")
                    if parsed["scm"].get("branches"):
                        message_parts.append(f"Branches: {', '.join(parsed['scm']['branches'])}")
                if parsed.get("triggers"):
                    message_parts.append("Triggers:")
                    for t in parsed["triggers"]:
                        message_parts.append(f"  - {t.get('type', '?')}: {t.get('schedule', 'N/A')}")
                if parsed.get("build_steps"):
                    message_parts.append("Build Steps:")
                    for s in parsed["build_steps"]:
                        message_parts.append(f"  - [{s.get('type', '?')}] {s.get('command', s.get('script', ''))[:200]}")
                if parsed.get("pipeline_script"):
                    message_parts.append(f"Pipeline Script:\n```groovy\n{parsed['pipeline_script'][:2000]}\n```")
                if parsed.get("post_build_actions"):
                    message_parts.append("Post-build Actions:")
                    for a in parsed["post_build_actions"]:
                        message_parts.append(f"  - {a.get('type', '?')}")
                if parsed.get("environment_vars"):
                    message_parts.append("Environment Variables:")
                    for ev in parsed["environment_vars"][:20]:
                        message_parts.append(f"  - {ev}")

                # Raw XML (truncated if very large)
                raw_truncated = raw_xml[:8000] if len(raw_xml) > 8000 else raw_xml
                message_parts.append(f"\nFull config.xml content:\n```xml\n{raw_truncated}\n```")

        # Process screenshot files
        if has_images:
            files = await db.files.find(
                {"id": {"$in": doc_create.file_ids}, "is_deleted": False},
                {"_id": 0}
            ).to_list(100)
            if files:
                message_parts.append("\nAdditionally, screenshots of the Jenkins pipeline are provided for visual context.")

        user_text = "\n".join(message_parts)

        # Generate documentation with AI
        chat = LlmChat(
            api_key=EMERGENT_KEY,
            session_id=f"doc-gen-{uuid.uuid4()}",
            system_message=base_prompt
        )

        if doc_create.ai_provider == "claude":
            chat.with_model("anthropic", "claude-sonnet-4-5-20250929")
        else:
            chat.with_model("gemini", "gemini-3-flash-preview")

        user_message = UserMessage(text=user_text)
        response = await chat.send_message(user_message)

        # Save documentation
        all_file_ids = doc_create.file_ids + doc_create.xml_file_ids
        doc_data = {
            "id": str(uuid.uuid4()),
            "title": doc_create.title,
            "content": response,
            "file_ids": all_file_ids,
            "template_id": doc_create.template_id,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        await db.documentations.insert_one(doc_data)

        return Documentation(**doc_data)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Generate error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/documentations", response_model=List[Documentation])
async def get_documentations():
    docs = await db.documentations.find({}, {"_id": 0}).sort("created_at", -1).to_list(100)
    return docs

@api_router.get("/documentations/{doc_id}", response_model=Documentation)
async def get_documentation(doc_id: str):
    doc = await db.documentations.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Documentation not found")
    return doc

@api_router.put("/documentations/{doc_id}", response_model=Documentation)
async def update_documentation(doc_id: str, update: DocumentationUpdate):
    doc = await db.documentations.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Documentation not found")
    
    update_data = update.model_dump(exclude_unset=True)
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.documentations.update_one({"id": doc_id}, {"$set": update_data})
    updated_doc = await db.documentations.find_one({"id": doc_id}, {"_id": 0})
    return updated_doc

@api_router.delete("/documentations/{doc_id}")
async def delete_documentation(doc_id: str):
    result = await db.documentations.delete_one({"id": doc_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Documentation not found")
    return {"message": "Documentation deleted"}

@api_router.post("/templates", response_model=Template)
async def create_template(template: TemplateCreate):
    template_data = {
        "id": str(uuid.uuid4()),
        **template.model_dump(),
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.templates.insert_one(template_data)
    return Template(**template_data)

@api_router.get("/templates", response_model=List[Template])
async def get_templates():
    templates = await db.templates.find({}, {"_id": 0}).to_list(100)
    return templates

@api_router.delete("/templates/{template_id}")
async def delete_template(template_id: str):
    result = await db.templates.delete_one({"id": template_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    return {"message": "Template deleted"}

@api_router.get("/export/pdf/{doc_id}")
async def export_pdf(doc_id: str):
    doc = await db.documentations.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Documentation not found")
    
    buffer = BytesIO()
    pdf_doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 47, 167),
        spaceAfter=30
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12
    )
    
    code_style = ParagraphStyle(
        'Code',
        parent=styles['Code'],
        fontSize=9,
        leftIndent=20,
        textColor=RGBColor(50, 50, 50),
        backColor=RGBColor(240, 240, 240)
    )
    
    story = []
    story.append(Paragraph(doc["title"], title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Parse markdown content
    lines = doc["content"].split('\n')
    for line in lines:
        if line.startswith('# '):
            story.append(Paragraph(line[2:], title_style))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], heading_style))
        elif line.startswith('```'):
            continue
        elif line.strip():
            story.append(Paragraph(line, styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
    
    pdf_doc.build(story)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{doc["title"]}.pdf"'}
    )

@api_router.get("/export/word/{doc_id}")
async def export_word(doc_id: str):
    doc = await db.documentations.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Documentation not found")
    
    word_doc = Document()
    
    # Add title
    title = word_doc.add_heading(doc["title"], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Parse and add content
    lines = doc["content"].split('\n')
    in_code_block = False
    
    for line in lines:
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if line.startswith('# '):
            word_doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            word_doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            word_doc.add_heading(line[4:], 3)
        elif in_code_block:
            p = word_doc.add_paragraph(line)
            p.style = 'Code'
        elif line.strip():
            word_doc.add_paragraph(line)
    
    buffer = BytesIO()
    word_doc.save(buffer)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{doc["title"]}.docx"'}
    )

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def startup():
    try:
        init_storage()
        logger.info("Storage initialized")
    except Exception as e:
        logger.error(f"Storage init failed: {e}")

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()